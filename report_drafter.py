"""
Report Drafter (Beta) — TCFD <-> TNFD first-draft generation.

Upload a TCFD report and generate a first draft of a TNFD report (or vice
versa). The drafter reuses the requirement mapping in
ReportingFrameworks_v1.xlsx: for each requirement of the target framework,
Claude searches the source report for transferable content and either

  - drafts disclosure prose grounded in that content (with page citations),
  - drafts what is supportable and flags what is missing, or
  - emits a gap box explaining what the company needs to add.

The drafter never invents target-domain content (e.g. nature data for a
TNFD draft) that is not present in the source report — anything without
supporting evidence in the report becomes a gap, not prose.
"""

import json

import anthropic
import streamlit as st
from io import BytesIO

# ============================================
# CONSTANTS
# ============================================

DRAFT_STATUS_DRAFTED = "Drafted"
DRAFT_STATUS_PARTIAL = "Partially drafted"
DRAFT_STATUS_GAP = "Gap"

DRAFT_STATUS_COLORS = {
    DRAFT_STATUS_DRAFTED: "#1C6B4A",
    DRAFT_STATUS_PARTIAL: "#C98A2B",
    DRAFT_STATUS_GAP: "#B4472F",
}

# Supported conversion directions (source framework, target framework)
DRAFT_DIRECTIONS = {
    "TCFD report → TNFD draft": ("TCFD", "TNFD"),
    "TNFD report → TCFD draft": ("TNFD", "TCFD"),
}

# Drafting models. Sonnet is recommended: drafting disclosure prose is a
# harder task than the classification done in the Report Analyser.
DRAFTER_MODELS = {
    "Claude Sonnet 4.5 — recommended for drafting quality": "claude-sonnet-4-5",
    "Claude Haiku 4.5 — lowest cost": "claude-haiku-4-5",
}

DOMAIN_NOTES = {
    # What must never be fabricated when drafting INTO this framework
    "TNFD": (
        "nature-related data: dependencies and impacts on nature, "
        "biodiversity metrics, ecosystem locations, LEAP assessment results, "
        "nature-related targets"
    ),
    "TCFD": (
        "climate-related data: GHG emissions figures, climate scenario "
        "analysis results, climate metrics and targets, transition plan "
        "details"
    ),
}


# ============================================
# DRAFT GENERATION (Claude)
# ============================================

def _build_system_message(source_fw, target_fw, source_full, target_full,
                          page_texts, first_page_number):
    """System message with the page-tagged source report — cached across
    the per-topic calls."""
    tagged_pages = "\n\n".join(
        f"[PAGE {first_page_number + i}] {text}"
        for i, text in enumerate(page_texts)
    )
    return [
        {
            "type": "text",
            "text": (
                "You are an expert sustainability disclosure writer. You are "
                f"drafting a first draft of a {target_full} ({target_fw}) "
                f"report for a company, using ONLY the content of the "
                f"company's existing {source_full} ({source_fw}) report "
                "provided below.\n\n"
                "For EACH target-framework requirement you are given, you "
                "must:\n"
                "1. Search the ENTIRE source report for content that can "
                "legitimately support the target disclosure. Relevant "
                "content may be spread across multiple sections.\n"
                "2. If the report supports it, write formal disclosure "
                "prose (1-3 paragraphs, third person, in the terminology "
                "and style of the target framework) based ONLY on what the "
                "report actually says.\n"
                "3. Record the page numbers (from the [PAGE n] markers) of "
                "every passage you drew on.\n"
                "4. If content is missing, describe specifically what the "
                "company needs to add.\n\n"
                "STRICT GROUNDING RULES:\n"
                f"- NEVER invent {DOMAIN_NOTES[target_fw]}. If the source "
                "report does not contain it, it is a gap.\n"
                "- Governance and risk-management PROCESSES described in "
                "the source report may be presented as the basis for the "
                "corresponding target disclosure, but insert a bracketed "
                "note such as \"[TO CONFIRM: board oversight extends to "
                f"{target_fw}-relevant issues]\" and mark the item as "
                "partial rather than complete.\n"
                "- Every drafted sentence must be traceable to the source "
                "report. Do not pad drafts with generic boilerplate.\n\n"
                "SOURCE REPORT TEXT (page-tagged):\n"
                f"{tagged_pages}"
            ),
            "cache_control": {"type": "ephemeral"},
        }
    ]


def _build_topic_prompt(target_fw, target_full, topic, reqs_list,
                        requirement_refs):
    refs = requirement_refs or {}
    prompt = (
        f"Draft the **{topic}** section of the {target_full} ({target_fw}) "
        "report. Work through each requirement below.\n\n"
    )
    for idx, req in enumerate(reqs_list, start=1):
        ref = refs.get((target_fw, req), "")
        ref_tag = f" (Source: {ref})" if ref else ""
        prompt += f"{idx}.{ref_tag} {req}\n"
    prompt += (
        "\n\nRespond ONLY with a JSON array, one element per requirement, "
        "in the same order. Each element must have exactly these keys:\n"
        "{\n"
        ' "reference": "<the Source reference exactly as given, or empty string>",\n'
        ' "requirement": "<the requirement text>",\n'
        ' "status": "<one of: drafted | partial | gap>",\n'
        ' "draft_text": "<formal disclosure prose grounded in the source report; empty string if status is gap>",\n'
        ' "source_pages": [<page numbers used, e.g. 12, 13>],\n'
        ' "gap_guidance": "<for partial: what must be added to complete this disclosure; for gap: 2-4 specific sentences on what the company needs to produce; empty string if status is drafted>"\n'
        "}\n\n"
        "Status meanings:\n"
        '- "drafted": the source report fully supports this disclosure.\n'
        '- "partial": some transferable content exists; draft what is '
        "supportable and use gap_guidance for the rest.\n"
        '- "gap": nothing in the source report supports this disclosure; '
        "draft_text must be an empty string.\n\n"
        "No markdown, no backticks, no preamble — just the raw JSON array."
    )
    return prompt


def claude_draft_report(page_texts, first_page_number, source_fw, target_fw,
                        api_key, model, framework_requirements,
                        requirement_refs, framework_full_names,
                        progress_bar=None, prettify_topic=None):
    """
    Generate a first-draft target-framework report from a source report.

    One API call per target-framework topic (pillar). The page-tagged
    source report sits in the cached system message, so cache reads keep
    the per-topic cost low. If a topic response is truncated or fails to
    parse, the topic is retried in two halves.

    Returns (draft_items, token_usage). Each draft item:
      {topic, reference, requirement, status, draft_text,
       source_pages, gap_guidance}
    """
    client = anthropic.Anthropic(api_key=api_key)

    source_full = framework_full_names.get(source_fw, source_fw)
    target_full = framework_full_names.get(target_fw, target_fw)
    system_message = _build_system_message(
        source_fw, target_fw, source_full, target_full,
        page_texts, first_page_number,
    )

    topics = framework_requirements.get(target_fw, {})
    prettify = prettify_topic or (lambda t: t)
    draft_items = []
    usage_totals = {
        "input_tokens": 0, "output_tokens": 0,
        "cache_read_tokens": 0, "cache_write_tokens": 0,
    }

    def _call_and_parse(prompt_text):
        resp = client.messages.create(
            model=model,
            max_tokens=16384,
            system=system_message,
            messages=[{"role": "user", "content": prompt_text}],
        )
        if resp.stop_reason == "max_tokens":
            raise ValueError("Response truncated (max_tokens reached)")
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
        items = json.loads(raw)
        usage_totals["input_tokens"] += resp.usage.input_tokens
        usage_totals["output_tokens"] += resp.usage.output_tokens
        usage_totals["cache_read_tokens"] += getattr(
            resp.usage, "cache_read_input_tokens", 0) or 0
        usage_totals["cache_write_tokens"] += getattr(
            resp.usage, "cache_creation_input_tokens", 0) or 0
        return items

    def _normalise(item, topic):
        raw_status = str(item.get("status", "gap")).strip().lower()
        if raw_status.startswith("draft"):
            status = DRAFT_STATUS_DRAFTED
        elif raw_status.startswith("part"):
            status = DRAFT_STATUS_PARTIAL
        else:
            status = DRAFT_STATUS_GAP

        draft_text = str(item.get("draft_text", "") or "").strip()
        pages = item.get("source_pages", []) or []
        pages = sorted({int(p) for p in pages if str(p).strip().isdigit()})

        # Enforce grounding: prose without page citations is not a
        # complete draft, and a "gap" never carries prose.
        if status == DRAFT_STATUS_DRAFTED and (not draft_text or not pages):
            status = DRAFT_STATUS_PARTIAL if draft_text else DRAFT_STATUS_GAP
        if status == DRAFT_STATUS_GAP:
            draft_text = ""

        return {
            "topic": topic,
            "reference": str(item.get("reference", "") or "").strip(),
            "requirement": str(item.get("requirement", "") or "").strip(),
            "status": status,
            "draft_text": draft_text,
            "source_pages": pages,
            "gap_guidance": str(item.get("gap_guidance", "") or "").strip(),
        }

    total_steps = max(len(topics), 1)
    for step, (raw_topic, reqs_list) in enumerate(topics.items()):
        topic = prettify(raw_topic)
        chunks = [reqs_list]
        items = None
        try:
            items = _call_and_parse(_build_topic_prompt(
                target_fw, target_full, topic, reqs_list, requirement_refs))
        except (ValueError, json.JSONDecodeError):
            # Truncated / unparseable — retry in two halves
            mid = max(len(reqs_list) // 2, 1)
            chunks = [reqs_list[:mid], reqs_list[mid:]]
        except anthropic.AuthenticationError:
            raise
        except anthropic.APIError as e:
            st.error(f"API error while drafting '{topic}': {e}")
            chunks = []

        if items is None and chunks and chunks != [reqs_list]:
            items = []
            for chunk in chunks:
                if not chunk:
                    continue
                try:
                    items.extend(_call_and_parse(_build_topic_prompt(
                        target_fw, target_full, topic, chunk,
                        requirement_refs)))
                except Exception as e:
                    st.warning(f"Could not draft part of '{topic}': {e}")

        for item in items or []:
            draft_items.append(_normalise(item, topic))

        if progress_bar:
            progress_bar.progress((step + 1) / total_steps)

    return draft_items, usage_totals


# ============================================
# WORD DOCUMENT EXPORT
# ============================================

def _shade_paragraph(paragraph, hex_fill):
    """Apply a background shade to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def generate_draft_docx(draft_items, source_fw, target_fw,
                        framework_full_names, source_name=""):
    """Assemble the draft report as a Word document. Returns BytesIO."""
    from datetime import date

    from docx import Document
    from docx.shared import Pt, RGBColor

    target_full = framework_full_names.get(target_fw, target_fw)
    source_full = framework_full_names.get(source_fw, source_fw)

    doc = Document()

    title = doc.add_heading(f"Draft {target_fw} Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2A)
    doc.add_paragraph(f"{target_full}")
    src_label = f" ({source_name})" if source_name else ""
    doc.add_paragraph(
        f"First draft generated from the company's {source_full} "
        f"({source_fw}) report{src_label} on {date.today().isoformat()}."
    )

    banner = doc.add_paragraph()
    run = banner.add_run(
        "BETA — AI-GENERATED FIRST DRAFT. This document was drafted "
        "automatically from the source report and is not professional "
        "advice. Every statement must be verified against the source "
        "report and company records before use. Highlighted boxes mark "
        "content that must be written by the company."
    )
    run.bold = True
    run.font.size = Pt(9)
    _shade_paragraph(banner, "F6E3C5")

    counts = {s: 0 for s in DRAFT_STATUS_COLORS}
    for item in draft_items:
        counts[item["status"]] += 1
    doc.add_paragraph(
        f"Coverage summary: {counts[DRAFT_STATUS_DRAFTED]} requirement(s) "
        f"drafted, {counts[DRAFT_STATUS_PARTIAL]} partially drafted, "
        f"{counts[DRAFT_STATUS_GAP]} gap(s) requiring new content."
    )

    current_topic = None
    for item in draft_items:
        if item["topic"] != current_topic:
            current_topic = item["topic"]
            doc.add_heading(current_topic, level=1)

        ref = f"{item['reference']} — " if item["reference"] else ""
        doc.add_heading(f"{ref}{item['requirement']}", level=2)

        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {item['status']}")
        status_run.bold = True
        status_run.font.size = Pt(9)

        if item["draft_text"]:
            for para_text in item["draft_text"].split("\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
            if item["source_pages"]:
                pages = ", ".join(str(p) for p in item["source_pages"])
                src = doc.add_paragraph()
                src_run = src.add_run(
                    f"[Drafted from {source_fw} report, page(s) {pages} — "
                    "verify before use]"
                )
                src_run.italic = True
                src_run.font.size = Pt(9)

        if item["status"] == DRAFT_STATUS_GAP:
            gap = doc.add_paragraph()
            gap_run = gap.add_run(
                "TO COMPLETE — no supporting content found in the source "
                f"report. {item['gap_guidance']}"
            )
            gap_run.bold = True
            _shade_paragraph(gap, "F2D3CB")
        elif item["status"] == DRAFT_STATUS_PARTIAL and item["gap_guidance"]:
            more = doc.add_paragraph()
            more_run = more.add_run(
                f"MORE NEEDED — {item['gap_guidance']}"
            )
            more_run.bold = True
            _shade_paragraph(more, "F6E3C5")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================
# STREAMLIT TAB
# ============================================

def render_drafter_tab(framework_requirements, requirement_refs,
                       framework_full_names, extract_text_from_pdf,
                       prettify_topic=None):
    """Render the Report Drafter (Beta) tab.

    Dependencies (requirement data, framework names, the PDF extractor,
    topic-name prettifier) are passed in from streamlit_app to avoid a
    circular import.
    """
    st.markdown(
        '<div style="display:inline-flex;align-items:center;gap:8px;'
        'background:#F6E3C5;border:1px solid #E3C893;border-radius:20px;'
        'padding:6px 14px;font-size:13px;font-weight:700;color:#8A5A12;'
        'margin-bottom:10px;">BETA</div>',
        unsafe_allow_html=True,
    )
    st.markdown("### Report Drafter")
    st.markdown(
        "Upload a **TCFD** report and generate a first draft of a **TNFD** "
        "report (or the reverse). The drafter writes only what your "
        "existing report supports — every drafted paragraph cites the "
        "source pages it came from, and anything your report doesn't "
        "cover becomes a clearly-marked gap explaining what you need to "
        "add. It will not invent nature or climate content that isn't "
        "there."
    )
    st.warning(
        "This is an AI-generated first draft, not professional advice. "
        "Verify every statement against your source report before use.",
        icon="⚠️",
    )
    st.markdown("---")

    ctrl_col, upload_col = st.columns([1, 1])

    with ctrl_col:
        secrets_key = st.secrets.get("ANTHROPIC_API_KEY", "")
        if secrets_key:
            api_key = secrets_key
            st.markdown(
                '<div style="background:#E8F2EA;border:1px solid #C6E0CC;'
                'border-radius:8px;padding:10px;font-size:13px;'
                'color:#1C6B4A;">✓ API key configured</div>',
                unsafe_allow_html=True,
            )
        else:
            api_key = st.text_input(
                "Anthropic API Key", type="password",
                placeholder="sk-ant-...", key="drafter_api_key",
                help=(
                    "Required for drafting. Your key is not stored. "
                    "Get one at console.anthropic.com"
                ),
            )

        direction_label = st.selectbox(
            "Conversion", list(DRAFT_DIRECTIONS.keys()),
            key="drafter_direction",
        )
        source_fw, target_fw = DRAFT_DIRECTIONS[direction_label]

        model_label = st.selectbox(
            "Drafting model", list(DRAFTER_MODELS.keys()),
            key="drafter_model",
        )
        model = DRAFTER_MODELS[model_label]

        n_reqs = sum(
            len(reqs)
            for reqs in framework_requirements.get(target_fw, {}).values()
        )
        n_topics = len(framework_requirements.get(target_fw, {}))
        st.markdown(
            f"Drafting against **{n_reqs}** {target_fw} requirements "
            f"across **{n_topics}** pillars "
            f"(~{n_topics} API calls; drafting produces long output, so "
            "expect this to cost more than an analysis run)."
        )

    with upload_col:
        st.markdown(f"**Upload your {source_fw} report**")
        uploaded_file = st.file_uploader(
            "Choose a PDF file", type="pdf", key="drafter_file",
            help=f"Upload the {source_fw} report to draft from",
        )

        page_start, page_end = 1, None
        if uploaded_file:
            import pymupdf
            pdf_bytes = uploaded_file.read()
            uploaded_file.seek(0)
            with pymupdf.open(stream=pdf_bytes, filetype="pdf") as doc:
                total_pages = len(doc)
            st.markdown(f"**PDF has {total_pages} pages.**")
            pr1, pr2 = st.columns(2)
            with pr1:
                page_start = st.number_input(
                    "From page", 1, total_pages, 1, key="drafter_ps")
            with pr2:
                page_end = st.number_input(
                    "To page", 1, total_pages, total_pages,
                    key="drafter_pe")
            if page_start > page_end:
                st.warning("'From page' must be ≤ 'To page'.")

    draft_disabled = not uploaded_file or not api_key

    if st.button(
        f"Generate {target_fw} Draft", disabled=draft_disabled,
        type="primary", key="drafter_go",
    ):
        if page_end is not None and page_start > page_end:
            st.error("'From page' must be ≤ 'To page'.")
            st.stop()

        with st.spinner("Extracting text from PDF..."):
            try:
                text_list = extract_text_from_pdf(uploaded_file)
                start_idx = max(0, page_start - 1)
                end_idx = (
                    page_end if page_end is not None else len(text_list)
                )
                text_list = text_list[start_idx:end_idx]
            except Exception as e:
                st.error(f"Failed to extract PDF: {e}")
                st.stop()

        st.markdown(f"### Drafting {target_fw} report with Claude...")
        progress_bar = st.progress(0)

        try:
            draft_items, token_usage = claude_draft_report(
                text_list, page_start, source_fw, target_fw,
                api_key, model, framework_requirements,
                requirement_refs, framework_full_names, progress_bar,
                prettify_topic,
            )
            st.session_state.draft_items = draft_items
            st.session_state.draft_token_usage = token_usage
            st.session_state.draft_direction = (source_fw, target_fw)
            st.session_state.draft_source_name = uploaded_file.name
            st.success("Draft complete!")
        except anthropic.AuthenticationError:
            st.error("Invalid API key. Please check your Anthropic API key.")
        except Exception as e:
            st.error(f"Drafting failed: {e}")

    # --- Results ---
    st.markdown("---")
    if not st.session_state.get("draft_items"):
        return

    draft_items = st.session_state.draft_items
    source_fw, target_fw = st.session_state.draft_direction
    token_usage = st.session_state.get("draft_token_usage", {})

    counts = {s: 0 for s in DRAFT_STATUS_COLORS}
    for item in draft_items:
        counts[item["status"]] += 1
    total = max(len(draft_items), 1)
    draft_score = (
        counts[DRAFT_STATUS_DRAFTED] + 0.5 * counts[DRAFT_STATUS_PARTIAL]
    ) / total

    st.markdown(f"### Draft {target_fw} Report")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Draft coverage", f"{draft_score:.0%}")
    m2.metric("Drafted", counts[DRAFT_STATUS_DRAFTED])
    m3.metric("Partially drafted", counts[DRAFT_STATUS_PARTIAL])
    m4.metric("Gaps to complete", counts[DRAFT_STATUS_GAP])

    docx_buf = generate_draft_docx(
        draft_items, source_fw, target_fw, framework_full_names,
        st.session_state.get("draft_source_name", ""),
    )
    st.download_button(
        f"Download draft {target_fw} report (.docx)",
        data=docx_buf,
        file_name=f"draft_{target_fw.lower()}_report.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        type="primary",
    )

    # Group by topic for display
    topics_seen = []
    for item in draft_items:
        if item["topic"] not in topics_seen:
            topics_seen.append(item["topic"])

    for topic in topics_seen:
        topic_items = [i for i in draft_items if i["topic"] == topic]
        topic_gaps = sum(
            1 for i in topic_items if i["status"] == DRAFT_STATUS_GAP
        )
        gap_note = f" · {topic_gaps} gap(s)" if topic_gaps else ""
        with st.expander(
            f"{topic} ({len(topic_items)} requirements{gap_note})",
            expanded=False,
        ):
            for item in topic_items:
                color = DRAFT_STATUS_COLORS[item["status"]]
                ref = (
                    f"<span style='font-family:monospace;font-size:12px;"
                    f"color:#8A9488;'>{item['reference']}</span> "
                    if item["reference"] else ""
                )
                st.markdown(
                    f"{ref}<strong>{item['requirement']}</strong><br>"
                    f"<span style='background:{color};color:#FCFAF3;"
                    f"border-radius:6px;padding:2px 10px;font-size:12px;"
                    f"font-weight:600;'>{item['status']}</span>",
                    unsafe_allow_html=True,
                )
                if item["draft_text"]:
                    st.markdown(
                        f"> {item['draft_text']}".replace("\n", "\n> ")
                    )
                    if item["source_pages"]:
                        pages = ", ".join(
                            str(p) for p in item["source_pages"])
                        st.caption(
                            f"Drafted from {source_fw} report, "
                            f"page(s) {pages} — verify before use."
                        )
                if item["status"] == DRAFT_STATUS_GAP:
                    st.error(
                        f"**To complete:** {item['gap_guidance']}",
                        icon="✍️",
                    )
                elif (
                    item["status"] == DRAFT_STATUS_PARTIAL
                    and item["gap_guidance"]
                ):
                    st.warning(
                        f"**More needed:** {item['gap_guidance']}",
                        icon="✍️",
                    )
                st.markdown("---")

    if token_usage:
        with st.expander("Token usage", expanded=False):
            st.markdown(
                f"- Input tokens: {token_usage.get('input_tokens', 0):,}\n"
                f"- Output tokens: {token_usage.get('output_tokens', 0):,}\n"
                f"- Cache read tokens: "
                f"{token_usage.get('cache_read_tokens', 0):,}\n"
                f"- Cache write tokens: "
                f"{token_usage.get('cache_write_tokens', 0):,}"
            )
